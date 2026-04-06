"""Tests for raoc.agents.execution.ExecutionAgent.

All tests use tmp_path and monkeypatch config.WORKSPACE / config.BACKUPS_DIR.
CommandWrapper is mocked; HostSampler is mocked or real.
"""

import shutil
from pathlib import Path
from unittest.mock import MagicMock

import pytest

from raoc import config
from raoc.agents.execution import ExecutionAgent
from raoc.db.queries import create_job, get_job
from raoc.db.schema import create_tables, get_engine
from raoc.models.action import ActionObject, ActionType
from raoc.models.job import JobStatus
from raoc.substrate.command_wrapper import CommandWrapper
from raoc.substrate.exceptions import UnsupportedFileTypeError
from raoc.substrate.host_sampler import HostSampler


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture()
def db(tmp_path):
    engine = get_engine(db_path=tmp_path / "test_exec.db")
    create_tables(engine)
    return engine


@pytest.fixture()
def workspace(tmp_path, monkeypatch):
    """Redirect WORKSPACE and BACKUPS_DIR to tmp_path for every test."""
    ws = tmp_path / "raoc_workspace"
    bk = ws / ".backups"
    ws.mkdir()
    bk.mkdir()
    monkeypatch.setattr(config, "WORKSPACE", ws)
    monkeypatch.setattr(config, "BACKUPS_DIR", bk)
    return ws


def _job(db, raw: str = "test job") -> str:
    return create_job(raw, engine=db).job_id


def _action(job_id: str, step: int, atype: ActionType, target: str,
            command: str | None = None) -> ActionObject:
    return ActionObject(
        job_id=job_id,
        step_index=step,
        action_type=atype,
        risk_level="low",
        target_path=target,
        intent="test",
        command=command,
    )


def _mock_wrapper(exit_code: int = 0, stdout: str = "ok", stderr: str = "") -> MagicMock:
    w = MagicMock(spec=CommandWrapper)
    w.run.return_value = {
        "exit_code": exit_code, "stdout": stdout,
        "stderr": stderr, "timed_out": False, "duration_ms": 10,
    }
    return w


def _mock_sampler(content: str = "safe content") -> MagicMock:
    s = MagicMock(spec=HostSampler)
    s.read_text_file.return_value = content
    return s


def _agent(db, workspace, wrapper=None, sampler=None) -> ExecutionAgent:
    return ExecutionAgent(
        db=db,
        command_wrapper=wrapper or _mock_wrapper(),
        sampler=sampler or _mock_sampler(),
    )


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

class TestFileBackup:
    """FILE_BACKUP creates the correct .bak file."""

    def test_bak_file_created(self, db, workspace):
        target = workspace / "notes.txt"
        target.write_text("original content")

        job_id = _job(db)
        action = _action(job_id, 0, ActionType.FILE_BACKUP, str(target))

        _agent(db, workspace).run(job_id, [action])

        # Backup is now timestamped — find it by glob
        baks = list(config.BACKUPS_DIR.glob("notes_*.txt.bak"))
        assert len(baks) == 1
        assert baks[0].read_text() == "original content"

    def test_backup_output_contains_path(self, db, workspace):
        target = workspace / "notes.txt"
        target.write_text("content")

        job_id = _job(db)
        action = _action(job_id, 0, ActionType.FILE_BACKUP, str(target))

        summary = _agent(db, workspace).run(job_id, [action])

        # Output should contain the timestamped backup path
        assert ".txt.bak" in summary["actions"][0]["output"]


class TestFileWrite:
    """FILE_WRITE writes content and handles failures."""

    def test_writes_correct_content(self, db, workspace):
        target = workspace / "notes.txt"
        target.write_text("old")

        job_id = _job(db)
        action = _action(job_id, 0, ActionType.FILE_WRITE, str(target), command="new content")

        _agent(db, workspace).run(job_id, [action])

        assert target.read_text() == "new content"

    def test_status_succeeded(self, db, workspace):
        target = workspace / "notes.txt"
        target.write_text("old")

        job_id = _job(db)
        action = _action(job_id, 0, ActionType.FILE_WRITE, str(target), command="new")

        summary = _agent(db, workspace).run(job_id, [action])

        assert summary["actions"][0]["status"] == "succeeded"

    def test_creates_parent_directory_if_missing(self, db, workspace):
        # Simulate writing a generated script to a scripts/ subdir that doesn't exist yet
        scripts_dir = workspace / "scripts"
        target = scripts_dir / "word_count.py"
        # scripts_dir intentionally NOT created beforehand

        job_id = _job(db)
        action = _action(job_id, 0, ActionType.FILE_WRITE, str(target), command="print('hello')")

        summary = _agent(db, workspace).run(job_id, [action])

        assert summary["actions"][0]["status"] == "succeeded"
        assert target.exists()
        assert target.read_text() == "print('hello')"

    def test_write_failure_restores_from_bak(self, db, workspace):
        from raoc.config import make_timestamped_stem
        from raoc.db.queries import get_job

        target = workspace / "notes.txt"
        target.write_text("original")

        job_id = _job(db)
        # Compute the timestamped backup name the agent will look for
        created_at = get_job(job_id, engine=db).created_at
        ts_stem = make_timestamped_stem(target.name, created_at)
        bak = config.BACKUPS_DIR / f"{ts_stem}{target.suffix}.bak"
        shutil.copy2(target, bak)

        # Make target read-only to force write failure
        target.chmod(0o444)

        action = _action(job_id, 0, ActionType.FILE_WRITE, str(target), command="should fail")

        try:
            summary = _agent(db, workspace).run(job_id, [action])
            assert summary["actions"][0]["status"] == "failed"
            # Restore should have run — content still "original"
            target.chmod(0o644)
            assert target.read_text() == "original"
        finally:
            target.chmod(0o644)


class TestCmdExecute:
    """CMD_EXECUTE captures stdout and stops on failure."""

    def test_captures_stdout(self, db, workspace):
        wrapper = _mock_wrapper(stdout="script output")
        job_id = _job(db)
        target = str(workspace / "script.py")
        action = _action(job_id, 0, ActionType.CMD_EXECUTE, target, command="python3 script.py")

        summary = _agent(db, workspace, wrapper=wrapper).run(job_id, [action])

        assert "script output" in summary["actions"][0]["output"]
        assert summary["actions"][0]["status"] == "succeeded"

    def test_nonzero_exit_code_stops_subsequent_steps(self, db, workspace):
        wrapper = _mock_wrapper(exit_code=1)
        job_id = _job(db)
        target = str(workspace / "script.py")

        actions = [
            _action(job_id, 0, ActionType.CMD_EXECUTE, target, command="python3 script.py"),
            _action(job_id, 1, ActionType.CMD_INSPECT, target),
        ]

        summary = _agent(db, workspace, wrapper=wrapper).run(job_id, actions)

        assert summary["steps_failed"] == 1
        # Only first action executed
        assert len(summary["actions"]) == 1
        assert summary["actions"][0]["status"] == "failed"

    def test_failed_cmd_sets_job_failed(self, db, workspace):
        wrapper = _mock_wrapper(exit_code=1)
        job_id = _job(db)
        target = str(workspace / "script.py")
        action = _action(job_id, 0, ActionType.CMD_EXECUTE, target, command="python3 script.py")

        _agent(db, workspace, wrapper=wrapper).run(job_id, [action])

        assert get_job(job_id, engine=db).status == JobStatus.FAILED


class TestCmdInspectBlocked:
    """Blocked pattern in CMD_INSPECT stops execution."""

    def test_blocked_pattern_stops_execution(self, db, workspace):
        sampler = _mock_sampler(content="rm -rf /home")
        job_id = _job(db)
        target = str(workspace / "bad.py")

        actions = [
            _action(job_id, 0, ActionType.CMD_INSPECT, target, command="rm -rf /home"),
            _action(job_id, 1, ActionType.CMD_EXECUTE, target, command="python3 bad.py"),
        ]

        summary = _agent(db, workspace, sampler=sampler).run(job_id, actions)

        assert summary["steps_failed"] == 1
        assert len(summary["actions"]) == 1
        assert summary["actions"][0]["status"] == "blocked"

    def test_blocked_sets_job_blocked(self, db, workspace):
        sampler = _mock_sampler(content="sudo rm everything")
        job_id = _job(db)
        target = str(workspace / "bad.py")
        action = _action(job_id, 0, ActionType.CMD_INSPECT, target, command="sudo rm")

        _agent(db, workspace, sampler=sampler).run(job_id, [action])

        assert get_job(job_id, engine=db).status == JobStatus.BLOCKED


class TestAllStepsSuccess:
    """All steps succeeding → job status VERIFYING."""

    def test_job_status_is_verifying(self, db, workspace):
        job_id = _job(db)
        action = _action(job_id, 0, ActionType.CMD_INSPECT, str(workspace))

        _agent(db, workspace).run(job_id, [action])

        assert get_job(job_id, engine=db).status == JobStatus.VERIFYING

    def test_summary_steps_completed(self, db, workspace):
        job_id = _job(db)
        actions = [
            _action(job_id, 0, ActionType.CMD_INSPECT, str(workspace)),
            _action(job_id, 1, ActionType.CMD_INSPECT, str(workspace)),
        ]

        summary = _agent(db, workspace).run(job_id, actions)

        assert summary["steps_completed"] == 2
        assert summary["steps_failed"] == 0


class TestFileWriteDocx:
    """FILE_WRITE to a .docx path creates a readable DOCX file."""

    def test_file_write_docx_creates_readable_docx(self, db, workspace):
        """Writing content via _do_file_write to a .docx path produces a valid DOCX."""
        from docx import Document

        target = workspace / "output.docx"
        content = "First line\nSecond line\nThird line"

        job_id = _job(db)
        action = _action(job_id, 0, ActionType.FILE_WRITE, str(target), command=content)
        action.detected_format = 'docx'

        summary = _agent(db, workspace).run(job_id, [action])

        assert summary["actions"][0]["status"] == "succeeded"
        assert target.exists()

        # Read back and verify paragraphs
        doc = Document(str(target))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "First line" in paragraphs
        assert "Second line" in paragraphs
        assert "Third line" in paragraphs

    def test_file_write_pdf_raises_unsupported(self, db, workspace):
        """_do_file_write with a .pdf target raises UnsupportedFileTypeError (via failed status)."""
        target = workspace / "output.pdf"
        content = "some content"

        job_id = _job(db)
        action = _action(job_id, 0, ActionType.FILE_WRITE, str(target), command=content)

        # The agent catches the exception and returns status=failed
        summary = _agent(db, workspace).run(job_id, [action])

        assert summary["actions"][0]["status"] == "failed"
        assert "not supported" in summary["actions"][0]["output"].lower() or \
               "unsupported" in summary["actions"][0]["output"].lower()

    def test_file_write_plain_text(self, db, workspace):
        """_do_file_write with a .txt target writes UTF-8 text content."""
        target = workspace / "notes.txt"
        content = "Plain text content for testing."

        job_id = _job(db)
        action = _action(job_id, 0, ActionType.FILE_WRITE, str(target), command=content)

        summary = _agent(db, workspace).run(job_id, [action])

        assert summary["actions"][0]["status"] == "succeeded"
        assert target.read_text(encoding='utf-8') == content

    def test_file_write_pdf_detected_format_raises(self, db, workspace):
        """_do_file_write with detected_format='pdf' returns failed status with clear message."""
        target = workspace / "output.txt"
        content = "some content"

        job_id = _job(db)
        action = _action(job_id, 0, ActionType.FILE_WRITE, str(target), command=content)
        action.detected_format = 'pdf'

        summary = _agent(db, workspace).run(job_id, [action])

        assert summary["actions"][0]["status"] == "failed"
        assert "not supported" in summary["actions"][0]["output"].lower() or \
               "unsupported" in summary["actions"][0]["output"].lower()


# ---------------------------------------------------------------------------
# Timestamp tests
# ---------------------------------------------------------------------------

class TestTimestampedFilenames:
    """Backup and PDF→DOCX output filenames must include a timestamp from job.created_at."""

    def test_backup_filename_includes_timestamp(self, db, workspace):
        """Backup created by _do_file_backup has the expected timestamped name."""
        from raoc.config import make_timestamped_stem
        from raoc.db.queries import get_job

        target = workspace / "notes.txt"
        target.write_text("original content")

        job_id = _job(db)
        action = _action(job_id, 0, ActionType.FILE_BACKUP, str(target))

        _agent(db, workspace).run(job_id, [action])

        created_at = get_job(job_id, engine=db).created_at
        ts_stem = make_timestamped_stem(target.name, created_at)
        expected_bak = config.BACKUPS_DIR / f"{ts_stem}{target.suffix}.bak"

        assert expected_bak.exists(), f"Expected backup at {expected_bak}"
        assert expected_bak.read_text() == "original content"

    def test_backup_filename_unique_across_jobs(self, db, workspace):
        """Two jobs on the same file create two distinct backup files."""
        from datetime import timezone

        target = workspace / "notes.txt"
        target.write_text("v1")

        # First job
        job_id_a = _job(db, "job a")
        action_a = _action(job_id_a, 0, ActionType.FILE_BACKUP, str(target))
        summary_a = _agent(db, workspace).run(job_id_a, [action_a])

        target.write_text("v2")

        # Second job — different created_at produces a different backup name
        import time
        time.sleep(1.05)  # ensure distinct second-level timestamp
        job_id_b = _job(db, "job b")
        action_b = _action(job_id_b, 0, ActionType.FILE_BACKUP, str(target))
        summary_b = _agent(db, workspace).run(job_id_b, [action_b])

        path_a = summary_a["actions"][0]["output"]
        path_b = summary_b["actions"][0]["output"]
        assert path_a != path_b, "Two jobs must produce different backup paths"

        from pathlib import Path
        assert Path(path_a).exists()
        assert Path(path_b).exists()

    def test_pdf_output_filename_includes_timestamp(self, db, workspace):
        """_do_file_write for a PDF → DOCX case writes to a timestamped .docx path."""
        from raoc.config import make_timestamped_stem
        from raoc.db.queries import get_job
        from docx import Document

        # Untimedstamped docx target — as if planning hadn't applied timestamp yet
        target = workspace / "proposal.docx"
        content = "Rewritten proposal content.\nSecond line."

        job_id = _job(db)
        action = _action(job_id, 0, ActionType.FILE_WRITE, str(target), command=content)
        action.detected_format = 'pdf'   # signals PDF→DOCX conversion

        _agent(db, workspace).run(job_id, [action])

        created_at = get_job(job_id, engine=db).created_at
        ts_stem = make_timestamped_stem(target.name, created_at)
        expected_output = workspace / f"{ts_stem}.docx"

        assert expected_output.exists(), f"Expected timestamped output at {expected_output}"
        doc = Document(str(expected_output))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "Rewritten proposal content." in paragraphs


# ---------------------------------------------------------------------------
# PDF in-place rewrite tests
# ---------------------------------------------------------------------------

class TestPdfInplace:
    """pdf_inplace write_strategy writes a modified .pdf; falls back to .docx on overflow."""

    def _make_pdf(self, path):
        """Create a minimal single-page text PDF at path."""
        from reportlab.pdfgen import canvas as rl_canvas
        c = rl_canvas.Canvas(str(path))
        c.drawString(72, 720, "Original text content in the PDF")
        c.save()

    def _text_blocks(self, page_num=0):
        """Return a minimal text_blocks list for testing (generous bbox)."""
        return [
            {
                "page": page_num,
                "bbox": (72, 500, 540, 750),  # generous area to avoid overflow
                "original_text": "Original text content in the PDF",
                "font_size": 12.0,
            }
        ]

    def test_pdf_inplace_writes_modified_pdf(self, db, workspace):
        """pdf_inplace strategy produces a timestamped .pdf output file."""
        from raoc.config import make_timestamped_stem
        from raoc.db.queries import get_job

        source_pdf = workspace / "report.pdf"
        self._make_pdf(source_pdf)

        job_id = _job(db)
        action = _action(job_id, 0, ActionType.FILE_WRITE, str(source_pdf),
                         command="OK")  # very short so it fits the block
        action.write_strategy = 'pdf_inplace'
        action.text_blocks = self._text_blocks()

        summary = _agent(db, workspace).run(job_id, [action])

        assert summary["actions"][0]["status"] == "succeeded"

        created_at = get_job(job_id, engine=db).created_at
        ts_stem = make_timestamped_stem(source_pdf.name, created_at)
        expected_pdf = workspace / f"{ts_stem}.pdf"
        assert expected_pdf.exists(), f"Expected timestamped PDF at {expected_pdf}"

    def test_pdf_inplace_fallback_on_overflow(self, db, workspace):
        """When insert_textbox overflows, execution falls back to .docx and logs fallback."""
        from unittest.mock import patch
        from raoc.db.queries import get_audit_log

        source_pdf = workspace / "report.pdf"
        self._make_pdf(source_pdf)

        job_id = _job(db)
        action = _action(job_id, 0, ActionType.FILE_WRITE, str(source_pdf),
                         command="Much longer rewritten text that won't fit in the original box")
        action.write_strategy = 'pdf_inplace'
        action.text_blocks = self._text_blocks()

        agent = _agent(db, workspace)

        # Simulate overflow: _pdf_insert_textbox returns negative value (overflow in pymupdf)
        with patch.object(agent, '_pdf_insert_textbox', return_value=-10.0):
            summary = agent.run(job_id, [action])

        assert summary["actions"][0]["status"] == "succeeded"

        # Output must be .docx, not .pdf
        output = summary["actions"][0]["output"] or ""
        assert "PDF_INPLACE_FALLBACK" in output

        # Audit log must contain fallback event
        log_events = [e["event"] for e in get_audit_log(job_id, engine=db)]
        assert "pdf_inplace_fallback" in log_events

        # A .docx file must exist; no .pdf output file
        from raoc.config import make_timestamped_stem
        from raoc.db.queries import get_job
        created_at = get_job(job_id, engine=db).created_at
        ts_stem = make_timestamped_stem(source_pdf.name, created_at)
        expected_docx = workspace / f"{ts_stem}.docx"
        expected_pdf = workspace / f"{ts_stem}.pdf"
        assert expected_docx.exists(), f"Expected fallback DOCX at {expected_docx}"
        assert not expected_pdf.exists(), "Should not have created .pdf on overflow fallback"

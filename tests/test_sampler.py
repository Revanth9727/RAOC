"""Tests for raoc.substrate.host_sampler — HostSampler."""

import pytest

from raoc.substrate.exceptions import (
    ExtractionError,
    FileTooLargeError,
    ScopeViolationError,
    ZipFileDetectedError,
)
from raoc.substrate.host_sampler import HostSampler


@pytest.fixture()
def sampler(tmp_path, monkeypatch):
    """Return a HostSampler with config.WORKSPACE pointed at tmp_path."""
    import raoc.config as cfg
    import raoc.substrate.host_sampler as hs_module

    monkeypatch.setattr(cfg, 'WORKSPACE', tmp_path)
    monkeypatch.setattr(hs_module, '_assert_in_workspace',
        lambda path: (
            (_ := None) or
            (_ := str(path.resolve()).startswith(str(tmp_path.resolve()))) or
            (_ := None)
        ) if str(path.resolve()).startswith(str(tmp_path.resolve())) else
        (_ for _ in ()).throw(ScopeViolationError(f"Path outside workspace: {path}"))
    )
    return HostSampler()


@pytest.fixture()
def workspace_sampler(tmp_path, monkeypatch):
    """Return (sampler, workspace_dir) with WORKSPACE = tmp_path."""
    import raoc.config as cfg

    monkeypatch.setattr(cfg, 'WORKSPACE', tmp_path)
    return HostSampler(), tmp_path


def test_sample_file_returns_correct_metadata(tmp_path, monkeypatch):
    """sample_file returns correct name, extension, and size_bytes."""
    import raoc.config as cfg
    monkeypatch.setattr(cfg, 'WORKSPACE', tmp_path)

    f = tmp_path / 'notes.txt'
    f.write_text('hello world')

    sampler = HostSampler()
    result = sampler.sample_file(f)

    assert result['name'] == 'notes.txt'
    assert result['extension'] == '.txt'
    assert result['size_bytes'] == len('hello world')
    assert result['exists'] is True
    assert result['is_locked'] is False


def test_sample_file_nonexistent_returns_exists_false(tmp_path, monkeypatch):
    """sample_file on a missing file returns exists=False without raising."""
    import raoc.config as cfg
    monkeypatch.setattr(cfg, 'WORKSPACE', tmp_path)

    sampler = HostSampler()
    result = sampler.sample_file(tmp_path / 'ghost.txt')

    assert result['exists'] is False
    assert result['size_bytes'] == 0


def test_sample_directory_returns_correct_file_count(tmp_path, monkeypatch):
    """sample_directory returns the number of files in the directory."""
    import raoc.config as cfg
    monkeypatch.setattr(cfg, 'WORKSPACE', tmp_path)

    (tmp_path / 'a.txt').write_text('aaa')
    (tmp_path / 'b.txt').write_text('bbb')
    (tmp_path / 'c.py').write_text('print()')

    sampler = HostSampler()
    result = sampler.sample_directory(tmp_path)

    assert result['file_count'] == 3
    assert result['total_size_bytes'] > 0
    assert len(result['files']) == 3


def test_read_text_file_returns_content(tmp_path, monkeypatch):
    """read_text_file returns the file content as a string."""
    import raoc.config as cfg
    monkeypatch.setattr(cfg, 'WORKSPACE', tmp_path)

    f = tmp_path / 'notes.txt'
    f.write_text('hello from raoc')

    sampler = HostSampler()
    content = sampler.read_text_file(f)

    assert content == 'hello from raoc'


def test_read_text_file_raises_for_oversized_file(tmp_path, monkeypatch):
    """read_text_file raises FileTooLargeError when content exceeds limit."""
    import raoc.config as cfg
    monkeypatch.setattr(cfg, 'WORKSPACE', tmp_path)
    monkeypatch.setattr(cfg, 'MAX_FILE_SIZE_CHARS', 10)

    f = tmp_path / 'big.txt'
    f.write_text('x' * 11)

    sampler = HostSampler()
    with pytest.raises(FileTooLargeError):
        sampler.read_text_file(f)


def test_sample_file_raises_scope_violation_outside_workspace(tmp_path, monkeypatch):
    """sample_file raises ScopeViolationError for a path outside WORKSPACE."""
    import raoc.config as cfg
    # Point WORKSPACE at a subdirectory so tmp_path itself is outside
    sub = tmp_path / 'workspace'
    sub.mkdir()
    monkeypatch.setattr(cfg, 'WORKSPACE', sub)

    outside = tmp_path / 'secret.txt'
    outside.write_text('forbidden')

    sampler = HostSampler()
    with pytest.raises(ScopeViolationError):
        sampler.sample_file(outside)


def test_is_text_file_always_returns_true(tmp_path):
    """is_text_file is deprecated and always returns True regardless of extension."""
    sampler = HostSampler()
    for ext in ('.txt', '.md', '.py', '.sh', '.csv', '.json', '.docx', '.pdf',
                '.png', '.jpg', '.zip', '.exe'):
        assert sampler.is_text_file(tmp_path / f'file{ext}') is True


def test_read_text_file_raises_for_unsupported_type(tmp_path, monkeypatch):
    """read_text_file raises ScopeViolationError for a non-text file type."""
    import raoc.config as cfg
    monkeypatch.setattr(cfg, 'WORKSPACE', tmp_path)

    f = tmp_path / 'image.png'
    f.write_bytes(b'\x89PNG\r\n')

    sampler = HostSampler()
    with pytest.raises(ScopeViolationError):
        sampler.read_text_file(f)


# ── Tests for read_docx and read_pdf_as_docx (utility methods) ────


def test_read_docx_extracts_text(tmp_path, monkeypatch):
    """read_docx returns paragraph text from a DOCX file."""
    import raoc.config as cfg
    from docx import Document

    monkeypatch.setattr(cfg, 'WORKSPACE', tmp_path)

    docx_path = tmp_path / 'notes.docx'
    doc = Document()
    doc.add_paragraph("Hello from DOCX")
    doc.add_paragraph("Second paragraph")
    doc.save(str(docx_path))

    sampler = HostSampler()
    text = sampler.read_docx(docx_path)

    assert "Hello from DOCX" in text
    assert "Second paragraph" in text


def test_read_pdf_as_docx_extracts_text(tmp_path, monkeypatch):
    """read_pdf_as_docx converts a PDF to DOCX and returns extracted text."""
    import raoc.config as cfg
    from reportlab.pdfgen import canvas as rl_canvas

    monkeypatch.setattr(cfg, 'WORKSPACE', tmp_path)

    pdf_path = tmp_path / 'report.pdf'
    c = rl_canvas.Canvas(str(pdf_path))
    c.drawString(100, 750, "Hello from PDF")
    c.save()

    docx_path = tmp_path / 'report.docx'
    sampler = HostSampler()
    text = sampler.read_pdf_as_docx(pdf_path, docx_path)

    assert isinstance(text, str)
    assert docx_path.exists()


# ── Tests for extract_text_for_rewrite ────────────────────────────


def test_extract_text_pdf_returns_text(tmp_path, monkeypatch):
    """extract_text_for_rewrite returns text, format='pdf', output_path ends in .docx."""
    import raoc.config as cfg
    from reportlab.pdfgen import canvas as rl_canvas

    monkeypatch.setattr(cfg, 'WORKSPACE', tmp_path)

    pdf_path = tmp_path / 'report.pdf'
    c = rl_canvas.Canvas(str(pdf_path))
    c.drawString(100, 750, "Hello from PDF extract")
    c.save()

    sampler = HostSampler()
    text, fmt, output_path, _blocks, _method = sampler.extract_text_for_rewrite(pdf_path)

    assert isinstance(text, str)
    assert fmt == 'pdf'
    assert output_path.suffix == '.docx'


def test_extract_text_docx_returns_text(tmp_path, monkeypatch):
    """extract_text_for_rewrite returns text, format='docx', output_path == input path."""
    import raoc.config as cfg
    from docx import Document

    monkeypatch.setattr(cfg, 'WORKSPACE', tmp_path)

    docx_path = tmp_path / 'notes.docx'
    doc = Document()
    doc.add_paragraph("DOCX content for extraction")
    doc.save(str(docx_path))

    sampler = HostSampler()
    text, fmt, output_path, _blocks, _method = sampler.extract_text_for_rewrite(docx_path)

    assert "DOCX content for extraction" in text
    assert fmt == 'docx'
    assert output_path == docx_path


def test_extract_text_plain_utf8(tmp_path, monkeypatch):
    """extract_text_for_rewrite returns text and format='text' for a plain UTF-8 file."""
    import raoc.config as cfg

    monkeypatch.setattr(cfg, 'WORKSPACE', tmp_path)

    txt_path = tmp_path / 'notes.txt'
    txt_path.write_text("Plain UTF-8 content", encoding='utf-8')

    sampler = HostSampler()
    text, fmt, output_path, _blocks, _method = sampler.extract_text_for_rewrite(txt_path)

    assert text == "Plain UTF-8 content"
    assert fmt == 'text'
    assert output_path == txt_path


def test_extract_text_plain_non_utf8_encoding(tmp_path, monkeypatch):
    """extract_text_for_rewrite succeeds for a latin-1 encoded file without error."""
    import raoc.config as cfg

    monkeypatch.setattr(cfg, 'WORKSPACE', tmp_path)

    latin1_path = tmp_path / 'latin1.txt'
    # Write latin-1 encoded content with accented characters
    content = "caf\xe9 na\xefve r\xe9sum\xe9 accented text"
    latin1_path.write_bytes(content.encode('latin-1'))

    sampler = HostSampler()
    text, fmt, output_path, _blocks, _method = sampler.extract_text_for_rewrite(latin1_path)

    assert isinstance(text, str)
    assert output_path == latin1_path


def test_extract_text_binary_raises_extraction_error(tmp_path, monkeypatch):
    """extract_text_for_rewrite raises ExtractionError for binary data."""
    import os
    import raoc.config as cfg

    monkeypatch.setattr(cfg, 'WORKSPACE', tmp_path)

    bin_path = tmp_path / 'random.bin'
    # Bytes 0x80-0x84 are UTF-8 continuation bytes; isolated they fail UTF-8 decode.
    # Bytes 0x01-0x03 are non-printable control chars (< 32, not tab/LF/CR).
    # Combined ratio of non-printable bytes exceeds MAX_BINARY_NONPRINTABLE_RATIO.
    bin_path.write_bytes(bytes([0x01, 0x02, 0x03, 0x80, 0x81, 0x82, 0x83, 0x84] * 25))

    sampler = HostSampler()
    with pytest.raises(ExtractionError) as exc_info:
        sampler.extract_text_for_rewrite(bin_path)

    assert "binary format" in str(exc_info.value).lower() or \
           "could not extract" in str(exc_info.value).lower()


def test_extract_text_too_large_raises_file_too_large(tmp_path, monkeypatch):
    """extract_text_for_rewrite raises FileTooLargeError when content exceeds limit."""
    import raoc.config as cfg

    monkeypatch.setattr(cfg, 'WORKSPACE', tmp_path)
    monkeypatch.setattr(cfg, 'MAX_FILE_SIZE_CHARS', 10)

    large_path = tmp_path / 'large.txt'
    large_path.write_text('x' * 11, encoding='utf-8')

    sampler = HostSampler()
    with pytest.raises(FileTooLargeError):
        sampler.extract_text_for_rewrite(large_path)


def test_detection_is_by_content_not_extension(tmp_path, monkeypatch):
    """A real PDF saved with a .txt extension is detected as pdf by content."""
    import raoc.config as cfg
    from reportlab.pdfgen import canvas as rl_canvas

    monkeypatch.setattr(cfg, 'WORKSPACE', tmp_path)

    # Create a real PDF but save it with .txt extension
    fake_txt = tmp_path / 'notreally.txt'
    c = rl_canvas.Canvas(str(fake_txt))
    c.drawString(100, 750, "Content detection test")
    c.save()

    sampler = HostSampler()
    text, fmt, output_path, _blocks, _method = sampler.extract_text_for_rewrite(fake_txt)

    # Format must be detected by content, not extension
    assert fmt == 'pdf'
    assert output_path.suffix == '.docx'


def test_pdf_output_path_is_docx(tmp_path, monkeypatch):
    """extract_text_for_rewrite on a PDF returns output_path with .docx suffix."""
    import raoc.config as cfg
    from reportlab.pdfgen import canvas as rl_canvas

    monkeypatch.setattr(cfg, 'WORKSPACE', tmp_path)

    pdf_path = tmp_path / 'notes.pdf'
    c = rl_canvas.Canvas(str(pdf_path))
    c.drawString(100, 750, "test")
    c.save()

    sampler = HostSampler()
    _text, _fmt, output_path, _blocks, _method = sampler.extract_text_for_rewrite(pdf_path)

    assert output_path.suffix == '.docx'


def test_binary_file_raises_extraction_error(tmp_path, monkeypatch):
    """extract_text_for_rewrite raises ExtractionError for a binary file (e.g. PNG)."""
    import raoc.config as cfg

    monkeypatch.setattr(cfg, 'WORKSPACE', tmp_path)

    # PNG magic bytes followed by many non-printable bytes
    png_path = tmp_path / 'image.png'
    png_path.write_bytes(b'\x89PNG\r\n\x1a\n' + bytes(range(256)) * 3)

    sampler = HostSampler()
    with pytest.raises(ExtractionError):
        sampler.extract_text_for_rewrite(png_path)


# ── ZIP file tests ────────────────────────────────────────────────


def test_zip_file_raises_zip_detected_error(tmp_path, monkeypatch):
    """extract_text_for_rewrite raises ZipFileDetectedError for a plain ZIP file."""
    import zipfile as zf_module
    import raoc.config as cfg

    monkeypatch.setattr(cfg, 'WORKSPACE', tmp_path)

    # Create a real ZIP with two text files
    zip_path = tmp_path / 'archive.zip'
    with zf_module.ZipFile(zip_path, 'w') as zf:
        zf.writestr('readme.txt', 'This is the readme.')
        zf.writestr('data.csv', 'col1,col2\n1,2\n')

    sampler = HostSampler()
    with pytest.raises(ZipFileDetectedError) as exc_info:
        sampler.extract_text_for_rewrite(zip_path)

    err = exc_info.value
    assert 'readme.txt' in err.contents or 'data.csv' in err.contents
    assert err.path == zip_path
    assert len(err.contents) == 2


def test_docx_not_mistaken_for_zip(tmp_path, monkeypatch):
    """extract_text_for_rewrite correctly identifies a DOCX file and does not raise ZipFileDetectedError."""
    import raoc.config as cfg
    from docx import Document

    monkeypatch.setattr(cfg, 'WORKSPACE', tmp_path)

    docx_path = tmp_path / 'notes.docx'
    doc = Document()
    doc.add_paragraph("Hello from DOCX")
    doc.save(str(docx_path))

    sampler = HostSampler()
    # Should NOT raise ZipFileDetectedError — DOCX is correctly handled
    text, fmt, output_path, _blocks, _method = sampler.extract_text_for_rewrite(docx_path)

    assert fmt == 'docx'
    assert 'Hello from DOCX' in text


# ── New PDF extraction tests ───────────────────────────────────────


def test_pdf_native_extraction_returns_blocks(tmp_path, monkeypatch):
    """extract_text_for_rewrite on a native-text PDF returns non-empty text_blocks."""
    import raoc.config as cfg
    from reportlab.pdfgen import canvas as rl_canvas

    monkeypatch.setattr(cfg, 'WORKSPACE', tmp_path)

    pdf_path = tmp_path / 'native.pdf'
    c = rl_canvas.Canvas(str(pdf_path))
    c.drawString(72, 720, "Hello from native PDF block")
    c.save()

    sampler = HostSampler()
    text, fmt, output_path, text_blocks, extraction_method = (
        sampler.extract_text_for_rewrite(pdf_path)
    )

    assert fmt == 'pdf'
    assert extraction_method == 'pdf_native'
    assert isinstance(text_blocks, list)
    assert len(text_blocks) > 0
    for block in text_blocks:
        assert 'page' in block
        assert 'bbox' in block
        assert 'original_text' in block
        assert 'font_size' in block


def test_pdf_ocr_fallback_on_empty_extraction(tmp_path, monkeypatch):
    """When pdfplumber returns empty text, OCR fallback is used."""
    import raoc.config as cfg
    from reportlab.pdfgen import canvas as rl_canvas
    from unittest.mock import patch, MagicMock

    monkeypatch.setattr(cfg, 'WORKSPACE', tmp_path)

    pdf_path = tmp_path / 'image_only.pdf'
    c = rl_canvas.Canvas(str(pdf_path))
    c.drawString(72, 720, "text that pdfplumber will not see")
    c.save()

    sampler = HostSampler()

    # Mock pdfplumber to return empty text (simulating image-only PDF)
    mock_page = MagicMock()
    mock_page.extract_text.return_value = ''
    mock_pdf = MagicMock()
    mock_pdf.__enter__ = MagicMock(return_value=mock_pdf)
    mock_pdf.__exit__ = MagicMock(return_value=False)
    mock_pdf.pages = [mock_page]

    # Mock pytesseract to return OCR text
    with patch('pdfplumber.open', return_value=mock_pdf):
        with patch('pytesseract.image_to_string', return_value='OCR extracted text'):
            text, fmt, output_path, text_blocks, extraction_method = (
                sampler.extract_text_for_rewrite(pdf_path)
            )

    assert extraction_method == 'pdf_ocr'
    assert 'OCR extracted text' in text
    assert text_blocks == []  # no bbox data for OCR path


def test_pdf_extraction_error_on_complete_failure(tmp_path, monkeypatch):
    """When both pdfplumber and pytesseract return empty, ExtractionError is raised."""
    import raoc.config as cfg
    from reportlab.pdfgen import canvas as rl_canvas
    from unittest.mock import patch, MagicMock

    monkeypatch.setattr(cfg, 'WORKSPACE', tmp_path)

    pdf_path = tmp_path / 'unreadable.pdf'
    c = rl_canvas.Canvas(str(pdf_path))
    c.save()

    sampler = HostSampler()

    mock_page = MagicMock()
    mock_page.extract_text.return_value = ''
    mock_pdf = MagicMock()
    mock_pdf.__enter__ = MagicMock(return_value=mock_pdf)
    mock_pdf.__exit__ = MagicMock(return_value=False)
    mock_pdf.pages = [mock_page]

    with patch('pdfplumber.open', return_value=mock_pdf):
        with patch('pytesseract.image_to_string', return_value='   '):
            with pytest.raises(ExtractionError) as exc_info:
                sampler.extract_text_for_rewrite(pdf_path)

    assert 'image-only' in str(exc_info.value).lower() or \
           'could not extract' in str(exc_info.value).lower()

"""File system and process state inspection for RAOC.

HostSampler is the only component that reads from the file system.
All methods enforce the workspace boundary via _assert_in_workspace().
"""

import logging
import zipfile
from datetime import datetime, timezone
from pathlib import Path

import psutil

from raoc import config
from raoc.substrate.exceptions import ScopeViolationError

logger = logging.getLogger(__name__)


def _assert_in_workspace(path: Path) -> None:
    """Raise ScopeViolationError if path is not inside config.WORKSPACE."""
    if not str(path.resolve()).startswith(str(config.WORKSPACE.resolve())):
        raise ScopeViolationError(f"Path outside workspace: {path}")


class HostSampler:
    """Reads file metadata, directory listings, and file content.

    All methods validate that the target path is inside config.WORKSPACE
    before performing any file system operation.
    """

    def is_locked(self, path: Path) -> bool:
        """Return True if any running process has this file open."""
        resolved = str(path.resolve())
        for proc in psutil.process_iter(['open_files']):
            try:
                for f in proc.info['open_files'] or []:
                    if f.path == resolved:
                        return True
            except (psutil.NoSuchProcess, psutil.AccessDenied):
                continue
        return False

    def sample_file(self, path: Path) -> dict:
        """Return metadata for a single file within the workspace.

        Returns a dict with path, name, extension, size_bytes, modified_at,
        created_at, exists, and is_locked.
        Raises ScopeViolationError if path is outside config.WORKSPACE.
        """
        _assert_in_workspace(path)
        exists = path.exists()
        if exists:
            stat = path.stat()
            size_bytes = stat.st_size
            modified_at = datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).isoformat()
            created_at = datetime.fromtimestamp(stat.st_ctime, tz=timezone.utc).isoformat()
            locked = self.is_locked(path)
        else:
            size_bytes = 0
            modified_at = None
            created_at = None
            locked = False

        return {
            'path':        str(path.resolve()),
            'name':        path.name,
            'extension':   path.suffix,
            'size_bytes':  size_bytes,
            'modified_at': modified_at,
            'created_at':  created_at,
            'exists':      exists,
            'is_locked':   locked,
        }

    def sample_directory(self, path: Path) -> dict:
        """Return a summary of a directory's immediate children.

        Returns dict with path, file_count, total_size_bytes, and files
        (list of sample_file results for each immediate child file).
        Non-recursive. Raises ScopeViolationError if outside workspace.
        """
        _assert_in_workspace(path)
        files = []
        total_size = 0
        for child in sorted(path.iterdir()):
            if child.is_file():
                meta = self.sample_file(child)
                files.append(meta)
                total_size += meta['size_bytes']
        return {
            'path':             str(path.resolve()),
            'file_count':       len(files),
            'total_size_bytes': total_size,
            'files':            files,
        }

    def is_text_file(self, path: Path) -> bool:
        """Deprecated: always returns True.

        Extension-based gating has been replaced by extract_text_for_rewrite(),
        which detects file type by content inspection. This method is kept for
        backward compatibility but should not be used for new code.
        """
        return True

    def read_docx(self, path: Path) -> str:
        """Extract all paragraph text from a DOCX file using python-docx.

        Joins paragraphs with newline characters.
        Raises ScopeViolationError if path is outside config.WORKSPACE.
        Raises FileTooLargeError if extracted text exceeds config.MAX_FILE_SIZE_CHARS.
        """
        from docx import Document
        from raoc.substrate.exceptions import FileTooLargeError

        _assert_in_workspace(path)
        doc = Document(str(path))
        text = "\n".join(para.text for para in doc.paragraphs)
        if len(text) > config.MAX_FILE_SIZE_CHARS:
            raise FileTooLargeError(
                f"File exceeds {config.MAX_FILE_SIZE_CHARS} chars: {len(text)}"
            )
        return text

    def read_pdf_as_docx(self, path: Path, output_path: Path) -> str:
        """Convert a PDF to DOCX at output_path and return extracted text.

        Uses pdf2docx for conversion, then reads the resulting DOCX with read_docx().
        output_path must be inside config.WORKSPACE.

        Raises ScopeViolationError if either path is outside config.WORKSPACE.
        Raises FileTooLargeError if extracted text exceeds config.MAX_FILE_SIZE_CHARS.
        """
        from pdf2docx import Converter
        from raoc.substrate.exceptions import FileTooLargeError

        _assert_in_workspace(path)
        _assert_in_workspace(output_path)

        cv = Converter(str(path))
        cv.convert(str(output_path), start=0, end=None)
        cv.close()

        return self.read_docx(output_path)

    def read_text_file(self, path: Path) -> str:
        """Read and return the content of a text file within the workspace.

        Raises ScopeViolationError if path is outside config.WORKSPACE or
        if the file cannot be decoded as UTF-8.
        Raises FileTooLargeError if content exceeds config.MAX_FILE_SIZE_CHARS.
        """
        from raoc.substrate.exceptions import FileTooLargeError

        _assert_in_workspace(path)
        try:
            content = path.read_text(encoding='utf-8')
        except UnicodeDecodeError:
            raise ScopeViolationError(
                f"Not a text file (could not decode as UTF-8): {path.name}"
            )
        if len(content) > config.MAX_FILE_SIZE_CHARS:
            raise FileTooLargeError(
                f"File exceeds {config.MAX_FILE_SIZE_CHARS} chars: {len(content)}"
            )
        return content

    def extract_pdf_text(self, path: Path) -> tuple:
        """Extract text from a PDF file with OCR fallback for image-only PDFs.

        Returns (full_text, extraction_method, text_blocks) where:
          extraction_method: 'pdf_native' or 'pdf_ocr'
          text_blocks: list of dicts with page, bbox, original_text, font_size.
            Empty list for OCR path (no bbox data available).

        Raises ExtractionError if both pdfplumber and OCR produce no text.
        """
        import pdfplumber

        from raoc.substrate.exceptions import ExtractionError

        # Try pdfplumber first for native text extraction
        full_text = ''
        try:
            with pdfplumber.open(str(path)) as pdf:
                pages_text = [page.extract_text() or '' for page in pdf.pages]
            full_text = '\n'.join(pages_text)
        except Exception:
            full_text = ''

        if full_text.strip():
            # Native text found — extract block-level bbox info using pymupdf
            text_blocks: list = []
            try:
                import fitz  # pymupdf
                doc = fitz.open(str(path))
                for page_num, page in enumerate(doc):
                    page_dict = page.get_text("dict")
                    for block in page_dict.get("blocks", []):
                        if block.get("type") != 0:  # skip image blocks
                            continue
                        block_text = ""
                        font_size = 12.0
                        for line in block.get("lines", []):
                            for span in line.get("spans", []):
                                block_text += span.get("text", "")
                                font_size = span.get("size", font_size)
                            block_text += "\n"
                        block_text = block_text.rstrip("\n")
                        if block_text.strip():
                            text_blocks.append({
                                "page": page_num,
                                "bbox": tuple(block["bbox"]),
                                "original_text": block_text,
                                "font_size": font_size,
                            })
                doc.close()
            except Exception:
                text_blocks = []
            return full_text, 'pdf_native', text_blocks

        # pdfplumber returned empty — try OCR via pymupdf + pytesseract
        try:
            import io

            import fitz
            import pytesseract
            from PIL import Image

            doc = fitz.open(str(path))
            ocr_pages: list = []
            for page in doc:
                mat = fitz.Matrix(300 / 72, 300 / 72)  # 300 DPI
                pix = page.get_pixmap(matrix=mat)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                ocr_pages.append(pytesseract.image_to_string(img))
            doc.close()

            ocr_text = '\n'.join(ocr_pages)
            if ocr_text.strip():
                return ocr_text, 'pdf_ocr', []  # no bbox data for OCR path
        except Exception:
            pass

        raise ExtractionError(
            f"Could not extract text from {path.name}. The PDF appears to be "
            f"image-only and OCR could not read it. It may be handwritten, "
            f"low resolution, or in an unsupported language."
        )

    def extract_text_for_rewrite(self, path: Path) -> tuple:
        """Extract text from a file for rewriting using content-based detection.

        Detection is always by file content (magic bytes), never by extension.
        Extension is used only as a hint for the write-back strategy.

        Returns (extracted_text, detected_format, output_path, text_blocks,
                 extraction_method) where:
          detected_format is one of: 'pdf', 'docx', 'text', 'unknown'
          output_path is where execution will write the result:
            pdf   → path.with_suffix(config.PDF_OUTPUT_EXTENSION)
            all others → path
          text_blocks: list of dicts for PDF files (page, bbox, original_text,
            font_size). Empty list for non-PDF files.
          extraction_method: 'pdf_native', 'pdf_ocr', or 'text'

        Detection order:
          1. Magic bytes: %PDF → pdf; PK\x03\x04 + word/ entry → docx
          2. PDF: extract with pdfplumber, OCR fallback via pymupdf+pytesseract
          3. DOCX: extract with python-docx
          4. Text: attempt UTF-8, fallback to chardet
          5. Unknown: binary ratio check, decode with errors='replace'

        Raises ExtractionError if text cannot be extracted.
        Raises FileTooLargeError if extracted text exceeds MAX_FILE_SIZE_CHARS.
        Raises ScopeViolationError if path is outside config.WORKSPACE.
        """
        import chardet

        from raoc.substrate.exceptions import ExtractionError, FileTooLargeError

        _assert_in_workspace(path)

        # Step 1: read magic bytes for format detection
        raw = path.read_bytes()
        magic = raw[:8]

        # ── PDF path ──────────────────────────────────────────────────
        if magic[:4] == b'%PDF':
            text, extraction_method, text_blocks = self.extract_pdf_text(path)
            if len(text) > config.MAX_FILE_SIZE_CHARS:
                raise FileTooLargeError(
                    f"File exceeds {config.MAX_FILE_SIZE_CHARS} chars: {len(text)}"
                )
            output_path = path.with_suffix(config.PDF_OUTPUT_EXTENSION)
            return text, 'pdf', output_path, text_blocks, extraction_method

        # ── DOCX / ZIP path ───────────────────────────────────────────
        if magic[:4] == b'PK\x03\x04':
            is_docx = False
            contents: list = []
            try:
                with zipfile.ZipFile(path) as zf:
                    namelist = zf.namelist()
                    is_docx = any(n.startswith('word/') for n in namelist)
                    if not is_docx:
                        contents = [n for n in namelist if not n.endswith('/')]
            except Exception:
                pass

            if is_docx:
                try:
                    from docx import Document
                    doc = Document(str(path))
                    text = '\n'.join(para.text for para in doc.paragraphs)
                except Exception as exc:
                    raise ExtractionError(
                        f"Could not extract text from {path.name}. "
                        f"The Word document may be corrupted or password-protected. "
                        f"Error: {exc}"
                    )
                logger.info(
                    "DOCX extraction: path=%s paragraphs=%d content_len=%d",
                    path, len(doc.paragraphs), len(text),
                )
                if not text or len(text.strip()) == 0:
                    raise ExtractionError(
                        f"DOCX extraction returned no text for {path.name}. "
                        f"The file has {len(doc.paragraphs)} paragraph(s) but all are blank. "
                        f"It may contain only images, tables, or non-text content that "
                        f"cannot be extracted as plain text."
                    )
                if len(text) > config.MAX_FILE_SIZE_CHARS:
                    raise FileTooLargeError(
                        f"File exceeds {config.MAX_FILE_SIZE_CHARS} chars: {len(text)}"
                    )
                return text, 'docx', path, [], 'text'
            # PK magic but not DOCX → it is a plain ZIP
            from raoc.substrate.exceptions import ZipFileDetectedError
            raise ZipFileDetectedError(path, contents)

        # ── Text path ─────────────────────────────────────────────────
        try:
            text = raw.decode('utf-8')
            if len(text) > config.MAX_FILE_SIZE_CHARS:
                raise FileTooLargeError(
                    f"File exceeds {config.MAX_FILE_SIZE_CHARS} chars: {len(text)}"
                )
            return text, 'text', path, [], 'text'
        except UnicodeDecodeError:
            pass
        except FileTooLargeError:
            raise

        # Try chardet for non-UTF-8 text files
        try:
            result = chardet.detect(raw)
            conf = result.get('confidence') or 0
            enc = result.get('encoding')
            if conf >= 0.7 and enc:
                text = raw.decode(enc)
                if len(text) > config.MAX_FILE_SIZE_CHARS:
                    raise FileTooLargeError(
                        f"File exceeds {config.MAX_FILE_SIZE_CHARS} chars: {len(text)}"
                    )
                return text, 'text', path, [], 'text'
        except (UnicodeDecodeError, LookupError):
            pass
        except FileTooLargeError:
            raise

        # ── Unknown fallback ──────────────────────────────────────────
        # Check non-printable byte ratio on first 4 KB
        sample = raw[:4096]
        if sample:
            non_printable = sum(
                1 for b in sample if b < 32 and b not in (9, 10, 13)
            )
            ratio = non_printable / len(sample)
            if ratio > config.MAX_BINARY_NONPRINTABLE_RATIO:
                raise ExtractionError(
                    f"Could not extract text from {path.name}. "
                    f"The file appears to be a binary format that cannot "
                    f"be rewritten as text."
                )

        text = raw.decode('utf-8', errors='replace')
        if len(text) > config.MAX_FILE_SIZE_CHARS:
            raise FileTooLargeError(
                f"File exceeds {config.MAX_FILE_SIZE_CHARS} chars: {len(text)}"
            )
        return text, 'unknown', path, [], 'text'
